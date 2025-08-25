# Fichier : utils/file_utils.py
# VERSION FINALE - Corrige le bug "TypeError: tuple indices..." et les bugs précédents.

import openpyxl
from openpyxl.utils import get_column_letter
from openpyxl.styles import Font
from datetime import datetime
import re
import logging
import docx
import os
import uuid

from db.database import DatabaseManager
from core.conges.manager import CongeManager
from utils.config_loader import CONFIG
from utils.date_utils import format_date_for_display

def _perform_db_operation_with_manager(db_path, certificats_path, operation_callback):
    """Fonction utilitaire pour gérer la connexion/déconnexion DB dans un thread."""
    db = DatabaseManager(db_path)
    if not db.connect():
        raise ConnectionError("Impossible de se connecter à la base de données depuis le thread.")
    
    manager = CongeManager(db, certificats_dir=certificats_path)

    try:
        return operation_callback(manager)
    finally:
        db.close()

def export_agents_to_excel(db_path, certificats_path, save_path):
    """Exporte la liste des agents. Conçu pour être exécuté dans un thread."""
    def operation(manager):
        agents = manager.get_all_agents()
        if not agents:
            return "Aucun agent à exporter."
        
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Agents"
        
        annee_exercice = manager.get_annee_exercice()
        an_n, an_n1, an_n2 = annee_exercice, annee_exercice - 1, annee_exercice - 2
        headers = ["ID", "Nom", "Prénom", "PPR", "Cadre", 
                   f"Solde {an_n2}", f"Solde {an_n1}", f"Solde {an_n}", "Solde Total Actif"]
        ws.append(headers)
        
        header_font = Font(bold=True)
        for cell in ws[1]:
            cell.font = header_font

        for agent in agents:
            soldes_par_annee = {s.annee: s.solde for s in agent.soldes_annuels if s.statut == 'Actif'}
            solde_total = agent.get_solde_total_actif()
            
            ws.append([agent.id, agent.nom, agent.prenom, agent.ppr, agent.cadre, 
                       soldes_par_annee.get(an_n2, 0.0), soldes_par_annee.get(an_n1, 0.0), 
                       soldes_par_annee.get(an_n, 0.0), solde_total])

        for col_idx, col_cells in enumerate(ws.columns, 1):
            max_length = max(len(str(cell.value or "")) for cell in col_cells)
            ws.column_dimensions[get_column_letter(col_idx)].width = max_length + 2
        
        os.makedirs(os.path.dirname(save_path), exist_ok=True)
        wb.save(save_path)
        return f"Liste des agents exportée avec succès vers\n{save_path}"

    return _perform_db_operation_with_manager(db_path, certificats_path, operation)

def export_all_conges_to_excel(db_path, certificats_path, save_path):
    """Exporte la liste de tous les congés. Conçu pour être exécuté dans un thread."""
    def operation(manager):
        all_conges = manager.get_all_conges()
        if not all_conges:
            return "Aucun congé à exporter."
            
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Tous les Congés"
        headers = ["Nom Agent", "Prénom Agent", "PPR Agent", "Type Congé", "Début", "Fin", "Jours Pris", "Statut", "Justification", "Intérimaire"]
        ws.append(headers)
        header_font = Font(bold=True)
        for cell in ws[1]:
            cell.font = header_font
            
        all_agents = {agent.id: agent for agent in manager.get_all_agents()}
        for conge in all_conges:
            agent = all_agents.get(conge.agent_id)
            agent_nom, agent_prenom, agent_ppr = (agent.nom, agent.prenom, agent.ppr) if agent else ("Agent", "Supprimé", "")
            interim_info = ""
            if conge.interim_id:
                interim = all_agents.get(conge.interim_id)
                interim_info = f"{interim.nom} {interim.prenom}" if interim else "Agent Supprimé"
            row_data = [agent_nom, agent_prenom, agent_ppr, conge.type_conge, format_date_for_display(conge.date_debut), format_date_for_display(conge.date_fin), conge.jours_pris, conge.statut, conge.justif or "", interim_info]
            ws.append(row_data)
            
        for col_idx, col_cells in enumerate(ws.columns, 1):
            max_length = max(len(str(cell.value or "")) for cell in col_cells)
            ws.column_dimensions[get_column_letter(col_idx)].width = max_length + 2

        os.makedirs(os.path.dirname(save_path), exist_ok=True)
        wb.save(save_path)
        return f"Tous les congés ont été exportés avec succès vers\n{save_path}"

    return _perform_db_operation_with_manager(db_path, certificats_path, operation)

def import_agents_from_excel(db_path, certificats_path, source_path):
    """Importe des agents avec une logique de colonnes optionnelles."""
    def operation(manager):
        errors = []
        added_count, updated_count = 0, 0
        
        required_headers = CONFIG.get('ui', {}).get('agent_import_headers_required', ['nom', 'prenom'])
        cadres_valides = CONFIG['ui'].get('grades', [])
        default_cadre = cadres_valides[0] if cadres_valides else "Administrateur"
        
        wb = openpyxl.load_workbook(source_path)
        ws = wb.active
        
        header = [str(cell.value or '').lower().strip() for cell in ws[1]]
        if not all(h in header for h in required_headers):
            raise ValueError(f"Colonnes requises manquantes : {', '.join(required_headers)}")

        col_map = {name: i for i, name in enumerate(header)}
        
        manager.db.conn.execute('BEGIN TRANSACTION')
        try:
            for i, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
                if all(c is None for c in row):
                    continue
                try:
                    # --- LECTURE SÉCURISÉE DES DONNÉES ---
                    nom = str(row[col_map['nom']] or '').strip()
                    prenom = str(row[col_map['prenom']] or '').strip()
                    if not nom or not prenom:
                        raise ValueError("Nom et prénom sont obligatoires.")

                    # Lecture sécurisée des colonnes optionnelles
                    ppr_idx = col_map.get('ppr')
                    ppr = str(row[ppr_idx] or '').strip() if ppr_idx is not None else ""

                    cadre_idx = col_map.get('cadre') # Utilisation de 'cadre' au lieu de 'grade'
                    cadre = str(row[cadre_idx] or '').strip() if cadre_idx is not None else ""
                    # --- FIN DE LA LECTURE SÉCURISÉE ---

                    if not ppr:
                        ppr_suffix = str(uuid.uuid4())[:8]
                        ppr = f"{nom.upper()[:4]}_{ppr_suffix}"

                    if not cadre:
                        cadre = default_cadre
                    
                    if cadre not in cadres_valides:
                        raise ValueError(f"Cadre '{cadre}' invalide. Cadres valides : {', '.join(cadres_valides)}")

                    soldes = {}
                    for col_name, col_idx in col_map.items():
                        match = re.match(r'solde_(\d{4})', col_name)
                        if match and row[col_idx] is not None:
                            annee = int(match.group(1))
                            solde_val = float(str(row[col_idx]).replace(",", "."))
                            if solde_val < 0:
                                raise ValueError(f"Solde négatif pour l'année {annee}.")
                            soldes[annee] = solde_val
                    
                    agent_data = {
                        'nom': nom, 'prenom': prenom, 'ppr': ppr, 'cadre': cadre,
                        'soldes': soldes
                    }
                    
                    existing_agent_id = manager.db.execute_query("SELECT id FROM agents WHERE ppr=?", (ppr,), fetch="one")
                    if existing_agent_id:
                        agent_data['id'] = existing_agent_id[0]
                        manager.save_full_agent(agent_data, is_modification=True)
                        updated_count += 1
                    else:
                        manager.save_full_agent(agent_data, is_modification=False)
                        added_count += 1
                
                except Exception as ve:
                    logging.warning(f"Erreur d'import à la ligne {i}: {ve}", exc_info=True)
                    errors.append(f"Ligne {i}: {ve}")
            
            if errors:
                manager.db.conn.rollback()
                raise Exception("Importation annulée en raison d'erreurs:\n" + "\n".join(errors[:10]))
            else:
                manager.db.conn.commit()
                return f"Importation réussie !\n\n- Agents ajoutés : {added_count}\n- Agents mis à jour : {updated_count}"

        except Exception as e:
            manager.db.conn.rollback()
            raise e

    return _perform_db_operation_with_manager(db_path, certificats_path, operation)

def generate_decision_from_template(template_path, output_path, context):
    """Génère un document Word à partir d'un modèle en remplaçant les tags."""
    try:
        doc = docx.Document(template_path)
        os.makedirs(os.path.dirname(output_path), exist_ok=True)

        for p in doc.paragraphs:
            for key, value in context.items():
                if key in p.text:
                    inline = p.runs
                    for i in range(len(inline)):
                        if key in inline[i].text:
                            text = inline[i].text.replace(key, str(value))
                            inline[i].text = text

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for key, value in context.items():
                            if key in p.text:
                                inline = p.runs
                                for i in range(len(inline)):
                                    if key in inline[i].text:
                                        text = inline[i].text.replace(key, str(value))
                                        inline[i].text = text
        doc.save(output_path)
        return True
    except Exception as e:
        logging.error(f"Erreur lors de la génération du document : {e}", exc_info=True)
        raise e